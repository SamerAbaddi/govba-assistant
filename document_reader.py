from io import BytesIO

import fitz
from docx import Document


def read_uploaded_file(uploaded_file) -> str:
    """
    Extract text from an uploaded TXT, DOCX, or PDF file.

    Parameters
    ----------
    uploaded_file:
        A file received from Streamlit's file uploader.

    Returns
    -------
    str
        Extracted document text.

    Raises
    ------
    ValueError
        If the file type is unsupported or no readable text is found.
    """

    if uploaded_file is None:
        return ""

    file_name = uploaded_file.name.lower()
    file_bytes = uploaded_file.getvalue()

    if file_name.endswith(".txt"):
        text = read_txt(file_bytes)

    elif file_name.endswith(".docx"):
        text = read_docx(file_bytes)

    elif file_name.endswith(".pdf"):
        text = read_pdf(file_bytes)

    else:
        raise ValueError(
            "Unsupported file type. Please upload a TXT, DOCX, or PDF file."
        )

    cleaned_text = text.strip()

    if not cleaned_text:
        raise ValueError(
            "No readable text was found. The document may be empty "
            "or the PDF may contain scanned images."
        )

    return cleaned_text


def read_txt(file_bytes: bytes) -> str:
    """Decode a plain-text file."""

    try:
        return file_bytes.decode("utf-8")

    except UnicodeDecodeError:
        return file_bytes.decode("latin-1")


def read_docx(file_bytes: bytes) -> str:
    """Extract paragraphs and tables from a Word document."""

    document = Document(BytesIO(file_bytes))
    content = []

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text.strip()

        if paragraph_text:
            content.append(paragraph_text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            row_text = " | ".join(cell for cell in cells if cell)

            if row_text:
                content.append(row_text)

    return "\n".join(content)


def read_pdf(file_bytes: bytes) -> str:
    """Extract text from a text-based PDF document."""

    content = []

    with fitz.open(stream=file_bytes, filetype="pdf") as document:
        for page_number, page in enumerate(document, start=1):
            page_text = page.get_text("text").strip()

            if page_text:
                content.append(
                    f"--- Page {page_number} ---\n{page_text}"
                )

    return "\n\n".join(content)